import io
import zipfile
from pathlib import Path
from typing import Dict, List, Tuple

import fitz  # PyMuPDF
import pandas as pd
import streamlit as st
from PIL import Image
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from streamlit_sortables import sort_items

st.set_page_config(
    page_title="PDF Fácil",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# =====================================================
# HELPERS
# =====================================================

def base_name(filename: str) -> str:
    return Path(filename).stem


def with_ext(filename: str, new_ext: str) -> str:
    return f"{base_name(filename)}{new_ext}"


def with_suffix(filename: str, suffix: str, ext: str = None) -> str:
    stem = base_name(filename)
    if ext is None:
        ext = Path(filename).suffix or ""
    return f"{stem}{suffix}{ext}"


def chunk_list(items, size):
    for i in range(0, len(items), size):
        yield items[i:i + size]


def human_size(num_bytes: int) -> str:
    value = float(num_bytes)
    for unit in ["B", "KB", "MB", "GB"]:
        if value < 1024 or unit == "GB":
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{num_bytes} B"


def safe_rerun():
    st.rerun()


def set_tool(tool_key: str):
    st.session_state.tool = tool_key
    safe_rerun()


def reset_editor_states():
    keys_to_remove = [
        "merge_editor_df", "merge_source",
        "split_editor_df", "split_source",
        "reorg_editor_df", "reorg_source",
    ]
    for key in keys_to_remove:
        if key in st.session_state:
            del st.session_state[key]


def unlock_pdf(pdf_bytes: bytes, password: str) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        result = reader.decrypt(password)
        if result == 0:
            raise ValueError("Senha incorreta.")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def merge_pdfs_from_plan(page_plan: pd.DataFrame, file_bytes_map: Dict[str, bytes]) -> bytes:
    writer = PdfWriter()
    selected = page_plan.sort_values("ordem")
    selected = selected[selected["incluir"] == True]

    for _, row in selected.iterrows():
        reader = PdfReader(io.BytesIO(file_bytes_map[row["arquivo"]]))
        writer.add_page(reader.pages[int(row["pagina_pdf"]) - 1])

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def split_pdf(file_bytes: bytes, split_after: int) -> Tuple[bytes, bytes]:
    reader = PdfReader(io.BytesIO(file_bytes))
    writer1 = PdfWriter()
    writer2 = PdfWriter()

    for i in range(split_after):
        writer1.add_page(reader.pages[i])

    for i in range(split_after, len(reader.pages)):
        writer2.add_page(reader.pages[i])

    out1 = io.BytesIO()
    out2 = io.BytesIO()
    writer1.write(out1)
    writer2.write(out2)
    return out1.getvalue(), out2.getvalue()


def rebuild_pdf_from_plan(page_plan: pd.DataFrame, file_bytes_map: Dict[str, bytes]) -> bytes:
    return merge_pdfs_from_plan(page_plan, file_bytes_map)


def compress_pdf_bytes(pdf_bytes: bytes) -> bytes:
    """
    Compressão simples usando PyMuPDF.
    Costuma funcionar melhor do que apenas regravar com pypdf.
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        out = io.BytesIO()
        doc.save(
            out,
            garbage=4,
            deflate=True,
            clean=True,
            incremental=False,
        )
        doc.close()
        return out.getvalue()
    except Exception:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        out = io.BytesIO()
        writer.write(out)
        return out.getvalue()


def images_to_pdf(image_files) -> bytes:
    images = []
    for img in image_files:
        image = Image.open(img).convert("RGB")
        images.append(image)

    if not images:
        raise ValueError("Envie pelo menos uma imagem.")

    out = io.BytesIO()
    images[0].save(out, save_all=True, append_images=images[1:], format="PDF")
    return out.getvalue()


def pdf_to_jpg_zip(pdf_bytes: bytes, quality: int = 90) -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for i in range(len(doc)):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="JPEG", quality=quality)
            zip_file.writestr(f"pagina_{i+1}.jpg", img_buffer.getvalue())

    doc.close()
    return zip_buffer.getvalue()


def pdf_text_to_docx(pdf_bytes: bytes) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    doc = Document()

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if i > 0:
            doc.add_page_break()
        doc.add_heading(f"Página {i+1}", level=1)
        for bloco in text.split("\n"):
            bloco = bloco.strip()
            if bloco:
                doc.add_paragraph(bloco)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def wrap_text_for_pdf(text: str, max_chars: int = 95) -> List[str]:
    words = text.split()
    if not words:
        return [""]

    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = current + " " + word
        if len(candidate) <= max_chars:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def docx_to_simple_pdf(docx_bytes: bytes) -> bytes:
    document = Document(io.BytesIO(docx_bytes))
    out = io.BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    width, height = A4
    x = 50
    y = height - 50
    line_height = 16

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            y -= line_height
        else:
            chunks = wrap_text_for_pdf(text, max_chars=95)
            for chunk in chunks:
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(x, y, chunk)
                y -= line_height

    c.save()
    return out.getvalue()


def get_pdf_thumbnail(pdf_bytes: bytes, page_number: int, zoom: float = 0.58):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc.load_page(page_number)
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    doc.close()
    return img


def build_merge_editor(files):
    rows = []
    for f in files:
        reader = PdfReader(io.BytesIO(f.getvalue()))
        for p in range(len(reader.pages)):
            rows.append(
                {
                    "id": f"{f.name}__{p+1}",
                    "arquivo": f.name,
                    "pagina_pdf": p + 1,
                    "incluir": True,
                    "ordem": len(rows) + 1,
                    "rotulo": f"{f.name} - página {p+1}",
                }
            )
    return pd.DataFrame(rows)


def build_single_editor(file_name: str, file_bytes: bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    rows = []
    for p in range(len(reader.pages)):
        rows.append(
            {
                "id": f"{file_name}__{p+1}",
                "arquivo": file_name,
                "pagina_pdf": p + 1,
                "incluir": True,
                "ordem": p + 1,
                "rotulo": f"Página {p+1}",
            }
        )
    return pd.DataFrame(rows)


def build_reorganize_editor(base_pdf, extra_pdfs):
    rows = []
    all_files = [base_pdf] + extra_pdfs
    for f in all_files:
        reader = PdfReader(io.BytesIO(f.getvalue()))
        for p in range(len(reader.pages)):
            rows.append(
                {
                    "id": f"{f.name}__{p+1}",
                    "arquivo": f.name,
                    "pagina_pdf": p + 1,
                    "incluir": True,
                    "ordem": len(rows) + 1,
                    "rotulo": f"{f.name} - página {p+1}",
                }
            )
    return pd.DataFrame(rows)


def apply_drag_order(df: pd.DataFrame, sorted_labels: List[str]) -> pd.DataFrame:
    label_to_row = {row["rotulo"]: row for _, row in df.iterrows()}
    ordered_rows = []
    for idx, label in enumerate(sorted_labels, start=1):
        row = label_to_row[label].copy()
        row["ordem"] = idx
        ordered_rows.append(row)
    return pd.DataFrame(ordered_rows)


def render_professional_preview(plan_df: pd.DataFrame, file_bytes_map: Dict[str, bytes], title="Preview das páginas"):
    selected = plan_df[plan_df["incluir"] == True].sort_values("ordem")

    st.markdown(f"### {title}")
    if selected.empty:
        st.info("Nenhuma página selecionada.")
        return

    records = selected.to_dict("records")
    for chunk in chunk_list(records, 4):
        cols = st.columns(4, gap="medium")
        for col, item in zip(cols, chunk):
            with col:
                try:
                    img = get_pdf_thumbnail(
                        file_bytes_map[item["arquivo"]],
                        int(item["pagina_pdf"]) - 1,
                        zoom=0.58,
                    )
                    st.markdown(
                        f"""
                        <div class="preview-page-card">
                            <div class="page-order-badge">{int(item["ordem"])}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    st.image(img, use_container_width=True)
                    st.markdown(
                        f"""
                        <div class="page-meta">
                            <strong>{item["arquivo"]}</strong><br>
                            Página {int(item["pagina_pdf"])}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                except Exception:
                    st.markdown(
                        f"""
                        <div class="preview-card-fallback">
                            Prévia indisponível<br>
                            {item["arquivo"]} - página {int(item["pagina_pdf"])}
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )


def render_image_preview(image_files):
    if not image_files:
        return

    st.markdown("### Preview das imagens")
    previews = list(image_files)[:8]
    for chunk in chunk_list(previews, 4):
        cols = st.columns(4, gap="medium")
        for col, img_file in zip(cols, chunk):
            with col:
                try:
                    img = Image.open(img_file)
                    st.image(img, use_container_width=True)
                    st.caption(img_file.name)
                except Exception:
                    st.caption(img_file.name)


def file_summary_box(files, accepted_label: str):
    if not files:
        return

    if not isinstance(files, list):
        files = [files]

    total_size = sum(getattr(f, "size", 0) for f in files)
    st.markdown(
        f"""
        <div class="upload-summary">
            <div><strong>{len(files)}</strong> arquivo(s) enviado(s)</div>
            <div>{accepted_label}</div>
            <div>Tamanho total: <strong>{human_size(total_size)}</strong></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# =====================================================
# CONFIG APP
# =====================================================

MAIN_TOOLS = {
    "unlock": "🔓 Remover senha",
    "merge": "📎 Juntar PDFs",
    "split": "✂️ Dividir PDF",
    "reorganize": "🗂️ Reorganizar PDF",
    "compress": "🗜️ Comprimir PDF",
}

CONVERSION_TOOLS = {
    "imgpdf": "🖼️ JPG para PDF",
    "pdfjpg": "🖼️ PDF para JPG",
    "pdfword": "📄 PDF para Word",
    "wordpdf": "📄 Word para PDF",
}

ALL_TOOLS = {}
ALL_TOOLS.update(MAIN_TOOLS)
ALL_TOOLS.update(CONVERSION_TOOLS)

DESCRIPTIONS = {
    "unlock": "Remova a senha de arquivos protegidos e baixe o PDF desbloqueado.",
    "merge": "Combine vários PDFs em um único arquivo e reorganize as páginas.",
    "split": "Separe páginas ou gere uma nova versão reorganizada do PDF.",
    "reorganize": "Reordene, exclua ou adicione páginas de outros PDFs.",
    "compress": "Reduza o tamanho do PDF com uma compressão simples e rápida.",
    "imgpdf": "Converta imagens JPG ou PNG em um único arquivo PDF.",
    "pdfjpg": "Transforme as páginas do PDF em imagens JPG em ZIP.",
    "pdfword": "Extraia o texto do PDF e gere um arquivo Word.",
    "wordpdf": "Converta DOCX em PDF simples com foco em texto.",
}

BANNERS = {
    "unlock": (
        "Remover senha de PDF",
        "Envie um ou mais PDFs protegidos, informe a senha e baixe os arquivos desbloqueados.",
    ),
    "merge": (
        "Juntar PDFs",
        "Envie vários PDFs, arraste as páginas para a ordem desejada e baixe tudo em um único arquivo.",
    ),
    "split": (
        "Dividir PDF",
        "Separe um PDF em partes ou reorganize páginas antes de baixar uma nova versão.",
    ),
    "reorganize": (
        "Reorganizar PDF",
        "Reordene, exclua e até adicione páginas de outros PDFs em um único documento final.",
    ),
    "compress": (
        "Comprimir PDF",
        "Reduza o tamanho do arquivo para facilitar envio, upload e compartilhamento.",
    ),
    "imgpdf": (
        "JPG para PDF",
        "Envie imagens, organize e gere um PDF de forma rápida.",
    ),
    "pdfjpg": (
        "PDF para JPG",
        "Converta páginas do PDF em imagens JPG prontas para baixar em ZIP.",
    ),
    "pdfword": (
        "PDF para Word",
        "Extraia o texto do PDF e crie um documento .docx.",
    ),
    "wordpdf": (
        "Word para PDF",
        "Converta seu arquivo .docx em PDF simples, ideal para textos.",
    ),
}

FEATURES = [
    "Sem instalação",
    "Funciona no celular",
    "Processamento rápido",
    "Várias ferramentas em um só lugar",
]

if "tool" not in st.session_state:
    st.session_state.tool = None

# =====================================================
# CSS - ESTILO MAIS ILOVEPDF
# =====================================================

st.markdown(
    """
<style>
:root {
    --bg: #f5f5f7;
    --card: #ffffff;
    --text: #1f2937;
    --muted: #6b7280;
    --line: #e5e7eb;
    --red: #e5322d;
    --red-dark: #cc2420;
    --soft: #fafafa;
    --green: #16a34a;
    --shadow: 0 10px 28px rgba(0,0,0,.05);
    --radius: 18px;
}

html, body, [data-testid="stAppViewContainer"] {
    background: var(--bg) !important;
    color: var(--text) !important;
}

header, [data-testid="stHeader"] {
    background: transparent !important;
}

#MainMenu {
    visibility: hidden;
}

footer {
    visibility: hidden;
}

[data-testid="stSidebar"] {
    display: none !important;
}

.block-container {
    padding-top: 1.5rem !important;
    padding-bottom: 2rem !important;
    max-width: 1320px !important;
}

.hero-wrap {
    margin: 0 auto 20px auto;
    max-width: 1180px;
}

.topbar {
    background: var(--card);
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 12px 16px;
    box-shadow: var(--shadow);
    margin-bottom: 18px;
}

.brand-row {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 16px;
    margin-bottom: 18px;
    flex-wrap: wrap;
}

.brand-left {
    display: flex;
    align-items: center;
    gap: 12px;
}

.brand-badge {
    width: 46px;
    height: 46px;
    border-radius: 14px;
    background: var(--red);
    display: inline-flex;
    align-items: center;
    justify-content: center;
    font-size: 24px;
}

.brand-title {
    margin: 0;
    font-size: 1.65rem;
    font-weight: 800;
    color: var(--text) !important;
}

.brand-subtitle {
    margin: 2px 0 0 0;
    color: var(--muted) !important;
    font-size: 0.96rem;
}

.hero-card {
    background: linear-gradient(180deg, #ffffff 0%, #fafafa 100%);
    border: 1px solid var(--line);
    border-radius: 24px;
    padding: 34px 28px;
    text-align: center;
    box-shadow: var(--shadow);
    margin-bottom: 22px;
}

.hero-card h1 {
    margin: 0 0 10px 0;
    font-size: 2.65rem;
    line-height: 1.1;
    color: var(--text) !important;
    font-weight: 900;
}

.hero-card p {
    max-width: 760px;
    margin: 0 auto;
    color: var(--muted) !important;
    font-size: 1.03rem;
}

.feature-chips {
    margin-top: 18px;
    display: flex;
    justify-content: center;
    gap: 10px;
    flex-wrap: wrap;
}

.feature-chip {
    background: #fff;
    border: 1px solid var(--line);
    color: var(--text);
    padding: 8px 12px;
    border-radius: 999px;
    font-size: 0.9rem;
    font-weight: 600;
}

.home-grid-wrap {
    max-width: 1180px;
    margin: 0 auto;
}

.tool-card {
    background: var(--card);
    border: 1px solid var(--line);
    border-radius: 22px;
    padding: 22px 20px;
    min-height: 220px;
    box-shadow: var(--shadow);
    transition: transform .18s ease, box-shadow .18s ease, border-color .18s ease;
    display: flex;
    flex-direction: column;
    justify-content: space-between;
}

.tool-card:hover {
    transform: translateY(-3px);
    border-color: #d1d5db;
    box-shadow: 0 14px 34px rgba(0,0,0,.08);
}

.tool-icon {
    width: 54px;
    height: 54px;
    border-radius: 16px;
    background: #fff1f1;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    font-size: 24px;
    margin-bottom: 16px;
}

.tool-title {
    font-size: 1.08rem;
    font-weight: 800;
    color: var(--text) !important;
    margin-bottom: 8px;
}

.tool-desc {
    font-size: 0.95rem;
    line-height: 1.5;
    color: var(--muted) !important;
    margin-bottom: 20px;
}

.tool-banner {
    max-width: 1180px;
    margin: 0 auto 16px auto;
    background: var(--card);
    border: 1px solid var(--line);
    border-radius: 24px;
    padding: 24px 22px;
    text-align: center;
    box-shadow: var(--shadow);
}

.tool-banner h2 {
    margin: 0;
    font-size: 2rem;
    font-weight: 900;
    color: var(--text) !important;
}

.tool-banner p {
    margin: 8px 0 0 0;
    color: var(--muted) !important;
    font-size: 1rem;
}

.tool-panel {
    max-width: 1180px;
    margin: 0 auto;
    background: var(--card);
    border: 1px solid var(--line);
    border-radius: 24px;
    padding: 24px 22px 28px 22px;
    box-shadow: var(--shadow);
}

.panel-section {
    background: #fcfcfc;
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 18px;
}

.panel-section h3 {
    margin-top: 0;
}

.editor-box {
    background: #fcfcfc;
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 18px;
}

.upload-summary {
    display: flex;
    gap: 14px;
    flex-wrap: wrap;
    align-items: center;
    justify-content: center;
    background: #fff;
    border: 1px solid var(--line);
    border-radius: 16px;
    padding: 12px 14px;
    margin: 12px 0 8px 0;
    color: var(--text) !important;
}

.preview-page-card {
    text-align: center;
    margin-bottom: 8px;
}

.page-order-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    min-width: 48px;
    height: 48px;
    border-radius: 999px;
    background: var(--red);
    color: #fff !important;
    font-size: 1.18rem;
    font-weight: 800;
    box-shadow: 0 8px 18px rgba(229, 50, 45, .22);
}

.page-meta {
    text-align: center;
    font-size: 0.88rem;
    color: var(--text) !important;
    margin-top: 8px;
    margin-bottom: 16px;
    background: #fff;
    border: 1px solid var(--line);
    border-radius: 12px;
    padding: 8px;
    min-height: 62px;
}

.preview-card-fallback {
    background: #fff;
    border: 1px solid var(--line);
    border-radius: 12px;
    padding: 14px;
    text-align: center;
}

.stButton > button,
[data-testid="stDownloadButton"] > button {
    width: 100%;
    border-radius: 14px !important;
    padding: 0.82rem 1rem !important;
    font-weight: 800 !important;
    border: none !important;
    transition: all .18s ease !important;
}

.stButton > button {
    background: var(--red) !important;
    color: white !important;
    box-shadow: 0 10px 22px rgba(229, 50, 45, .18) !important;
}

.stButton > button:hover {
    background: var(--red-dark) !important;
    transform: translateY(-1px);
}

[data-testid="stDownloadButton"] > button {
    background: #111827 !important;
    color: white !important;
}

[data-testid="stDownloadButton"] > button:hover {
    background: #000 !important;
}

.nav-btn .stButton > button {
    min-height: 54px !important;
    background: #fff !important;
    color: var(--text) !important;
    border: 1px solid var(--line) !important;
    box-shadow: none !important;
}

.nav-btn .stButton > button:hover {
    border-color: #c7cbd1 !important;
    background: #fafafa !important;
}

.home-use-btn .stButton > button {
    background: var(--red) !important;
    color: #fff !important;
}

.back-btn .stButton > button {
    background: #fff !important;
    color: var(--text) !important;
    border: 1px solid var(--line) !important;
    box-shadow: none !important;
}

[data-testid="stFileUploader"] {
    background: transparent !important;
}

[data-testid="stFileUploaderDropzone"] {
    background: #fff !important;
    border: 2px dashed #d1d5db !important;
    border-radius: 18px !important;
    padding-top: 18px !important;
    padding-bottom: 18px !important;
}

[data-testid="stFileUploaderDropzone"]:hover {
    border-color: var(--red) !important;
    background: #fffafa !important;
}

[data-testid="stFileUploaderDropzoneInstructions"] div {
    visibility: hidden;
    position: relative;
}

[data-testid="stFileUploaderDropzoneInstructions"] div::before {
    content: "Arraste e solte os arquivos aqui";
    visibility: visible;
    position: absolute;
    inset: 0;
    text-align: center;
    color: var(--text) !important;
    font-weight: 700;
}

[data-testid="stFileUploader"] section button {
    font-size: 0 !important;
    background: #fff5f5 !important;
    color: var(--red) !important;
    border: 1px solid #ffd4d4 !important;
    border-radius: 12px !important;
}

[data-testid="stFileUploader"] section button::after {
    content: "Selecionar arquivos";
    font-size: 14px !important;
    color: var(--red) !important;
    font-weight: 800;
}

label {
    color: var(--text) !important;
    font-weight: 700 !important;
}

input, textarea {
    color: var(--text) !important;
}

[data-baseweb="input"],
[data-baseweb="select"] > div {
    border-radius: 12px !important;
}

.footer-note {
    text-align: center;
    color: var(--muted) !important;
    margin-top: 26px;
    font-size: 0.96rem;
}

.small-muted {
    color: var(--muted) !important;
    font-size: 0.92rem;
}

hr {
    border: none;
    border-top: 1px solid var(--line);
    margin: 18px 0;
}

@media (max-width: 900px) {
    .hero-card {
        padding: 24px 16px;
    }

    .hero-card h1 {
        font-size: 2rem;
    }

    .tool-banner h2 {
        font-size: 1.6rem;
    }

    .tool-panel {
        padding: 16px 14px 22px 14px;
        border-radius: 18px;
    }

    .tool-card {
        min-height: 200px;
        padding: 18px 16px;
    }

    .upload-summary {
        justify-content: flex-start;
    }
}
</style>
""",
    unsafe_allow_html=True,
)

# =====================================================
# HEADER
# =====================================================

st.markdown(
    """
<div class="hero-wrap">
    <div class="topbar">
        <div class="brand-row">
            <div class="brand-left">
                <div class="brand-badge">📄</div>
                <div>
                    <div class="brand-title">PDF Fácil</div>
                    <div class="brand-subtitle">Ferramentas online para PDF com visual simples, rápido e profissional</div>
                </div>
            </div>
        </div>
    </div>

    <div class="hero-card">
        <h1>Edite, converta e organize PDFs em segundos</h1>
        <p>
            Remova senha, junte arquivos, divida páginas, comprima PDFs e faça conversões
            direto no navegador, inclusive no celular.
        </p>
        <div class="feature-chips">
            """
    + "".join([f'<div class="feature-chip">{item}</div>' for item in FEATURES]) +
    """
        </div>
    </div>
</div>
""",
    unsafe_allow_html=True,
)

# =====================================================
# TOP MENU
# =====================================================

st.markdown('<div class="home-grid-wrap">', unsafe_allow_html=True)

nav_cols = st.columns([1, 1, 1, 1, 1, 1.2], gap="small")

main_order = ["unlock", "merge", "split", "reorganize", "compress"]
for idx, key in enumerate(main_order):
    with nav_cols[idx]:
        st.markdown('<div class="nav-btn">', unsafe_allow_html=True)
        if st.button(MAIN_TOOLS[key], key=f"top_{key}"):
            reset_editor_states()
            set_tool(key)
        st.markdown("</div>", unsafe_allow_html=True)

with nav_cols[-1]:
    st.markdown('<div class="nav-btn">', unsafe_allow_html=True)
    with st.popover("Converter PDF"):
        for key, label in CONVERSION_TOOLS.items():
            if st.button(label, key=f"popover_{key}"):
                reset_editor_states()
                set_tool(key)
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# =====================================================
# HOME
# =====================================================

if st.session_state.tool is None:
    st.markdown('<div class="home-grid-wrap">', unsafe_allow_html=True)

    home_order = [
        "unlock", "merge", "split",
        "reorganize", "compress", "imgpdf",
        "pdfjpg", "pdfword", "wordpdf",
    ]

    for row in chunk_list(home_order, 3):
        cols = st.columns(3, gap="medium")
        for col, key in zip(cols, row):
            with col:
                emoji = ALL_TOOLS[key].split()[0]
                label = ALL_TOOLS[key]
                st.markdown(
                    f"""
                    <div class="tool-card">
                        <div>
                            <div class="tool-icon">{emoji}</div>
                            <div class="tool-title">{label}</div>
                            <div class="tool-desc">{DESCRIPTIONS[key]}</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                st.markdown('<div class="home-use-btn">', unsafe_allow_html=True)
                if st.button("Usar agora", key=f"home_{key}"):
                    reset_editor_states()
                    set_tool(key)
                st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='height:12px;'></div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

# =====================================================
# TOOL PANEL
# =====================================================

if st.session_state.tool is not None:
    st.markdown('<div class="home-grid-wrap">', unsafe_allow_html=True)

    back_cols = st.columns([1, 1.6, 1], gap="medium")
    with back_cols[1]:
        st.markdown('<div class="back-btn">', unsafe_allow_html=True)
        if st.button("⬅ Voltar para ferramentas", key="btn_back_home"):
            reset_editor_states()
            st.session_state.tool = None
            safe_rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    banner_title, banner_desc = BANNERS[st.session_state.tool]
    st.markdown(
        f"""
        <div class="tool-banner">
            <h2>{banner_title}</h2>
            <p>{banner_desc}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="tool-panel">', unsafe_allow_html=True)

    # =================================================
    # UNLOCK
    # =================================================
    if st.session_state.tool == "unlock":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivos = st.file_uploader(
            "Envie os PDFs protegidos",
            type=["pdf"],
            accept_multiple_files=True,
            key="unlock_files",
        )
        file_summary_box(arquivos, "PDF")
        senha = st.text_input("Digite a senha do PDF", type="password")
        st.markdown("</div>", unsafe_allow_html=True)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            desbloquear = st.button("Desbloquear PDFs", key="btn_unlock")

        if desbloquear:
            if not arquivos:
                st.warning("Envie pelo menos um PDF.")
            elif not senha:
                st.warning("Digite a senha.")
            else:
                try:
                    with st.spinner("Desbloqueando arquivos..."):
                        if len(arquivos) == 1:
                            unlocked = unlock_pdf(arquivos[0].getvalue(), senha)
                            st.success("PDF desbloqueado com sucesso.")
                            st.download_button(
                                "Baixar PDF desbloqueado",
                                unlocked,
                                arquivos[0].name,
                                key="download_unlock_single",
                            )
                        else:
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for arquivo in arquivos:
                                    unlocked = unlock_pdf(arquivo.getvalue(), senha)
                                    zip_file.writestr(arquivo.name, unlocked)

                            st.success("Arquivos desbloqueados com sucesso.")
                            st.download_button(
                                "Baixar todos em ZIP",
                                zip_buffer.getvalue(),
                                "pdfs_desbloqueados.zip",
                                key="download_unlock_all",
                            )
                except Exception as e:
                    st.error(str(e))

    # =================================================
    # MERGE
    # =================================================
    elif st.session_state.tool == "merge":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivos = st.file_uploader(
            "Selecione os PDFs",
            type=["pdf"],
            accept_multiple_files=True,
            key="merge_files",
        )
        file_summary_box(arquivos, "PDF")
        st.markdown("</div>", unsafe_allow_html=True)

        file_bytes_map = {}
        if arquivos:
            file_bytes_map = {a.name: a.getvalue() for a in arquivos}
            current_source = [a.name for a in arquivos]

            if "merge_editor_df" not in st.session_state or st.session_state.get("merge_source") != current_source:
                st.session_state.merge_editor_df = build_merge_editor(arquivos)
                st.session_state.merge_source = current_source

            st.markdown('<div class="editor-box">', unsafe_allow_html=True)
            st.markdown("### Ordem das páginas")
            st.caption("Arraste os itens abaixo para reorganizar a ordem final do PDF.")
            labels = st.session_state.merge_editor_df.sort_values("ordem")["rotulo"].tolist()
            sorted_labels = sort_items(labels, direction="horizontal", key="merge_sort_center")
            st.session_state.merge_editor_df = apply_drag_order(st.session_state.merge_editor_df, sorted_labels)
            st.markdown("</div>", unsafe_allow_html=True)

            with st.expander("Abrir editor detalhado de páginas", expanded=True):
                edited_df = st.data_editor(
                    st.session_state.merge_editor_df,
                    use_container_width=True,
                    hide_index=True,
                    num_rows="fixed",
                    column_config={
                        "id": None,
                        "arquivo": st.column_config.TextColumn("Arquivo", disabled=True),
                        "pagina_pdf": st.column_config.NumberColumn("Página", disabled=True),
                        "incluir": st.column_config.CheckboxColumn("Incluir"),
                        "ordem": st.column_config.NumberColumn("Ordem", disabled=True),
                        "rotulo": st.column_config.TextColumn("Rótulo", disabled=True),
                    },
                    key="merge_editor_grid",
                )
                st.session_state.merge_editor_df = edited_df

            render_professional_preview(
                st.session_state.merge_editor_df,
                file_bytes_map,
                "Preview das páginas selecionadas",
            )

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            juntar = st.button("Juntar PDFs", key="btn_merge")

        if juntar:
            if not arquivos:
                st.warning("Envie pelo menos um PDF.")
            else:
                selected = st.session_state.merge_editor_df
                if selected[selected["incluir"] == True].empty:
                    st.warning("Selecione pelo menos uma página.")
                else:
                    with st.spinner("Gerando PDF final..."):
                        merged = merge_pdfs_from_plan(st.session_state.merge_editor_df, file_bytes_map)
                    st.success("PDF unido com sucesso.")
                    st.download_button(
                        "Baixar PDF unido",
                        merged,
                        "pdf_unido.pdf",
                        key="download_merge",
                    )

    # =================================================
    # SPLIT
    # =================================================
    elif st.session_state.tool == "split":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivo = st.file_uploader(
            "Selecione o PDF",
            type=["pdf"],
            accept_multiple_files=False,
            key="split_file",
        )
        file_summary_box(arquivo, "PDF")
        st.markdown("</div>", unsafe_allow_html=True)

        if arquivo:
            file_bytes = arquivo.getvalue()
            file_map = {arquivo.name: file_bytes}
            reader = PdfReader(io.BytesIO(file_bytes))
            paginas = len(reader.pages)

            if "split_editor_df" not in st.session_state or st.session_state.get("split_source") != arquivo.name:
                st.session_state.split_editor_df = build_single_editor(arquivo.name, file_bytes)
                st.session_state.split_source = arquivo.name

            st.markdown('<div class="editor-box">', unsafe_allow_html=True)
            st.markdown("### Reorganizar páginas")
            st.caption("Arraste as páginas e escolha quais devem entrar no PDF final.")
            labels = st.session_state.split_editor_df.sort_values("ordem")["rotulo"].tolist()
            sorted_labels = sort_items(labels, direction="horizontal", key="split_sort_center")
            st.session_state.split_editor_df = apply_drag_order(st.session_state.split_editor_df, sorted_labels)
            st.markdown("</div>", unsafe_allow_html=True)

            with st.expander("Abrir editor detalhado de páginas", expanded=True):
                edited_df = st.data_editor(
                    st.session_state.split_editor_df,
                    use_container_width=True,
                    hide_index=True,
                    num_rows="fixed",
                    column_config={
                        "id": None,
                        "arquivo": st.column_config.TextColumn("Arquivo", disabled=True),
                        "pagina_pdf": st.column_config.NumberColumn("Página", disabled=True),
                        "incluir": st.column_config.CheckboxColumn("Incluir"),
                        "ordem": st.column_config.NumberColumn("Ordem", disabled=True),
                        "rotulo": st.column_config.TextColumn("Rótulo", disabled=True),
                    },
                    key="split_editor_grid",
                )
                st.session_state.split_editor_df = edited_df

            render_professional_preview(
                st.session_state.split_editor_df,
                file_map,
                "Preview das páginas",
            )

            st.markdown('<div class="panel-section">', unsafe_allow_html=True)
            st.markdown("### Divisão clássica")
            if paginas > 1:
                pagina = st.number_input(
                    "Dividir após a página",
                    min_value=1,
                    max_value=paginas - 1,
                    value=1,
                )

                c1, c2 = st.columns(2, gap="medium")
                with c1:
                    dividir = st.button("Dividir PDF", key="btn_split")
                with c2:
                    gerar_reorganizado = st.button("Gerar PDF reorganizado", key="btn_split_rebuild")

                if dividir:
                    with st.spinner("Dividindo PDF..."):
                        parte1, parte2 = split_pdf(file_bytes, pagina)
                    st.success("PDF dividido com sucesso.")
                    d1, d2 = st.columns(2, gap="medium")
                    with d1:
                        st.download_button(
                            "Baixar parte 1",
                            parte1,
                            with_suffix(arquivo.name, "_parte1", ".pdf"),
                            key="download_split_1",
                        )
                    with d2:
                        st.download_button(
                            "Baixar parte 2",
                            parte2,
                            with_suffix(arquivo.name, "_parte2", ".pdf"),
                            key="download_split_2",
                        )

                if gerar_reorganizado:
                    selected = st.session_state.split_editor_df
                    if selected[selected["incluir"] == True].empty:
                        st.warning("Selecione pelo menos uma página.")
                    else:
                        with st.spinner("Gerando PDF reorganizado..."):
                            rebuilt = rebuild_pdf_from_plan(st.session_state.split_editor_df, file_map)
                        st.success("PDF reorganizado gerado com sucesso.")
                        st.download_button(
                            "Baixar PDF reorganizado",
                            rebuilt,
                            with_suffix(arquivo.name, "_reorganizado", ".pdf"),
                            key="download_split_rebuild",
                        )
            else:
                st.info("Este PDF possui apenas uma página.")
            st.markdown("</div>", unsafe_allow_html=True)

    # =================================================
    # REORGANIZE
    # =================================================
    elif st.session_state.tool == "reorganize":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        base_pdf = st.file_uploader(
            "Selecione o PDF principal",
            type=["pdf"],
            accept_multiple_files=False,
            key="reorg_base",
        )
        extra_pdfs = st.file_uploader(
            "Selecione PDFs extras para adicionar páginas",
            type=["pdf"],
            accept_multiple_files=True,
            key="reorg_extra",
        )
        if base_pdf:
            resumo_lista = [base_pdf] + (extra_pdfs if extra_pdfs else [])
            file_summary_box(resumo_lista, "PDF")
        st.markdown("</div>", unsafe_allow_html=True)

        extras = extra_pdfs if extra_pdfs else []

        if base_pdf:
            all_names = [base_pdf.name] + [e.name for e in extras]
            file_bytes_map = {base_pdf.name: base_pdf.getvalue()}
            for e in extras:
                file_bytes_map[e.name] = e.getvalue()

            if "reorg_editor_df" not in st.session_state or st.session_state.get("reorg_source") != all_names:
                st.session_state.reorg_editor_df = build_reorganize_editor(base_pdf, extras)
                st.session_state.reorg_source = all_names

            st.markdown('<div class="editor-box">', unsafe_allow_html=True)
            st.markdown("### Reorganização final")
            st.caption("Arraste as páginas para definir a ordem do PDF final.")
            labels = st.session_state.reorg_editor_df.sort_values("ordem")["rotulo"].tolist()
            sorted_labels = sort_items(labels, direction="horizontal", key="reorg_sort_center")
            st.session_state.reorg_editor_df = apply_drag_order(st.session_state.reorg_editor_df, sorted_labels)
            st.markdown("</div>", unsafe_allow_html=True)

            with st.expander("Abrir editor detalhado de páginas", expanded=True):
                edited_df = st.data_editor(
                    st.session_state.reorg_editor_df,
                    use_container_width=True,
                    hide_index=True,
                    num_rows="fixed",
                    column_config={
                        "id": None,
                        "arquivo": st.column_config.TextColumn("Arquivo", disabled=True),
                        "pagina_pdf": st.column_config.NumberColumn("Página", disabled=True),
                        "incluir": st.column_config.CheckboxColumn("Incluir"),
                        "ordem": st.column_config.NumberColumn("Ordem", disabled=True),
                        "rotulo": st.column_config.TextColumn("Rótulo", disabled=True),
                    },
                    key="reorg_editor_grid",
                )
                st.session_state.reorg_editor_df = edited_df

            render_professional_preview(
                st.session_state.reorg_editor_df,
                file_bytes_map,
                "Preview do PDF reorganizado",
            )

            action_cols = st.columns([1, 1, 1], gap="medium")
            with action_cols[1]:
                gerar_reorg = st.button("Gerar PDF reorganizado", key="btn_reorg")

            if gerar_reorg:
                selected = st.session_state.reorg_editor_df
                if selected[selected["incluir"] == True].empty:
                    st.warning("Selecione pelo menos uma página.")
                else:
                    with st.spinner("Montando PDF final..."):
                        rebuilt = rebuild_pdf_from_plan(st.session_state.reorg_editor_df, file_bytes_map)
                    st.success("PDF reorganizado gerado com sucesso.")
                    st.download_button(
                        "Baixar PDF reorganizado",
                        rebuilt,
                        with_suffix(base_pdf.name, "_reorganizado", ".pdf"),
                        key="download_reorg",
                    )

    # =================================================
    # COMPRESS
    # =================================================
    elif st.session_state.tool == "compress":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivos = st.file_uploader(
            "Selecione os PDFs",
            type=["pdf"],
            accept_multiple_files=True,
            key="compress_files",
        )
        file_summary_box(arquivos, "PDF")
        st.caption("A redução varia de acordo com o conteúdo do arquivo. PDFs já otimizados podem reduzir pouco.")
        st.markdown("</div>", unsafe_allow_html=True)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            comprimir = st.button("Comprimir PDFs", key="btn_compress")

        if comprimir:
            if not arquivos:
                st.warning("Envie pelo menos um PDF.")
            else:
                try:
                    with st.spinner("Comprimindo arquivos..."):
                        if len(arquivos) == 1:
                            original_name = arquivos[0].name
                            original_bytes = arquivos[0].getvalue()
                            compressed = compress_pdf_bytes(original_bytes)
                            reduction = len(original_bytes) - len(compressed)

                            if reduction > 0:
                                st.success(
                                    f"Compressão concluída. Economia aproximada: {human_size(reduction)}."
                                )
                            else:
                                st.info("Compressão concluída. Este arquivo já estava bem otimizado.")

                            st.download_button(
                                "Baixar PDF comprimido",
                                compressed,
                                with_suffix(original_name, "_comprimido", ".pdf"),
                                key="download_compress_single",
                            )
                        else:
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for arquivo in arquivos:
                                    compressed = compress_pdf_bytes(arquivo.getvalue())
                                    zip_file.writestr(
                                        with_suffix(arquivo.name, "_comprimido", ".pdf"),
                                        compressed,
                                    )

                            st.success("Arquivos comprimidos com sucesso.")
                            st.download_button(
                                "Baixar PDFs comprimidos em ZIP",
                                zip_buffer.getvalue(),
                                "pdfs_comprimidos.zip",
                                key="download_compress",
                            )
                except Exception as e:
                    st.error(f"Não foi possível comprimir os arquivos: {e}")

    # =================================================
    # IMG -> PDF
    # =================================================
    elif st.session_state.tool == "imgpdf":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        imagens = st.file_uploader(
            "Envie imagens JPG ou PNG",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="imgpdf_files",
        )
        file_summary_box(imagens, "JPG / PNG")
        st.markdown("</div>", unsafe_allow_html=True)

        if imagens:
            render_image_preview(imagens)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            converter = st.button("Converter em PDF", key="btn_imgpdf")

        if converter:
            if not imagens:
                st.warning("Envie pelo menos uma imagem.")
            else:
                try:
                    with st.spinner("Convertendo imagens em PDF..."):
                        pdf_bytes = images_to_pdf(imagens)
                    output_name = with_ext(imagens[0].name, ".pdf") if len(imagens) == 1 else "imagens_convertidas.pdf"
                    st.success("PDF gerado com sucesso.")
                    st.download_button(
                        "Baixar PDF",
                        pdf_bytes,
                        output_name,
                        key="download_imgpdf",
                    )
                except Exception as e:
                    st.error(f"Não foi possível gerar o PDF: {e}")

    # =================================================
    # PDF -> JPG
    # =================================================
    elif st.session_state.tool == "pdfjpg":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivo = st.file_uploader(
            "Selecione o PDF",
            type=["pdf"],
            accept_multiple_files=False,
            key="pdfjpg_file",
        )
        file_summary_box(arquivo, "PDF")
        quality = st.slider("Qualidade do JPG", min_value=60, max_value=100, value=90, step=5)
        st.markdown("</div>", unsafe_allow_html=True)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            converter_jpg = st.button("Converter para JPG", key="btn_pdfjpg")

        if converter_jpg:
            if not arquivo:
                st.warning("Envie um PDF.")
            else:
                try:
                    with st.spinner("Convertendo PDF para JPG..."):
                        zip_bytes = pdf_to_jpg_zip(arquivo.getvalue(), quality=quality)
                    st.success("Conversão concluída com sucesso.")
                    st.download_button(
                        "Baixar JPGs em ZIP",
                        zip_bytes,
                        with_suffix(arquivo.name, "_jpg", ".zip"),
                        key="download_pdfjpg",
                    )
                except Exception as e:
                    st.error(f"Não foi possível converter este PDF: {e}")

    # =================================================
    # PDF -> WORD
    # =================================================
    elif st.session_state.tool == "pdfword":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivo = st.file_uploader(
            "Selecione o PDF",
            type=["pdf"],
            accept_multiple_files=False,
            key="pdfword_file",
        )
        file_summary_box(arquivo, "PDF")
        st.caption("Esta conversão extrai texto. PDFs digitalizados em imagem podem ter resultado limitado.")
        st.markdown("</div>", unsafe_allow_html=True)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            converter_word = st.button("Converter para Word", key="btn_pdfword")

        if converter_word:
            if not arquivo:
                st.warning("Envie um PDF.")
            else:
                try:
                    with st.spinner("Convertendo PDF para Word..."):
                        docx_bytes = pdf_text_to_docx(arquivo.getvalue())
                    st.success("Arquivo Word gerado com sucesso.")
                    st.download_button(
                        "Baixar Word (.docx)",
                        docx_bytes,
                        with_ext(arquivo.name, ".docx"),
                        key="download_pdfword",
                    )
                except Exception as e:
                    st.error(f"Não foi possível converter este PDF: {e}")

    # =================================================
    # WORD -> PDF
    # =================================================
    elif st.session_state.tool == "wordpdf":
        st.markdown('<div class="panel-section">', unsafe_allow_html=True)
        arquivo = st.file_uploader(
            "Selecione o arquivo Word (.docx)",
            type=["docx"],
            accept_multiple_files=False,
            key="wordpdf_file",
        )
        file_summary_box(arquivo, "DOCX")
        st.caption("Conversão simples, focada em texto.")
        st.markdown("</div>", unsafe_allow_html=True)

        action_cols = st.columns([1, 1, 1], gap="medium")
        with action_cols[1]:
            converter_pdf = st.button("Converter para PDF", key="btn_wordpdf")

        if converter_pdf:
            if not arquivo:
                st.warning("Envie um arquivo .docx.")
            else:
                try:
                    with st.spinner("Convertendo Word para PDF..."):
                        pdf_bytes = docx_to_simple_pdf(arquivo.getvalue())
                    st.success("PDF gerado com sucesso.")
                    st.download_button(
                        "Baixar PDF",
                        pdf_bytes,
                        with_ext(arquivo.name, ".pdf"),
                        key="download_wordpdf",
                    )
                except Exception as e:
                    st.error(f"Não foi possível converter este Word: {e}")

    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

st.markdown(
    '<div class="footer-note">PDF Fácil • Ferramentas gratuitas para PDF no navegador</div>',
    unsafe_allow_html=True,
)
